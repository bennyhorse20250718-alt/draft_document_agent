"""
Export service — convert draft text to DOCX or PDF.
"""
import os
from pathlib import Path
from typing import Literal
import logging

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import get_settings

logger = logging.getLogger(__name__)


class ExportService:
    def __init__(self):
        settings = get_settings()
        os.makedirs(settings.export_path, exist_ok=True)
        self._export_path = settings.export_path

    def export_docx(self, draft_text: str, filename: str = "draft") -> str:
        """Export draft text to DOCX. Returns the saved file path."""
        doc = Document()

        # Margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        for line in draft_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            # Simple heuristic: lines in ALL CAPS or starting with '#' are headings
            if stripped.startswith("# "):
                p = doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                p = doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                p = doc.add_heading(stripped[4:], level=3)
            elif stripped.isupper() and len(stripped) < 80:
                p = doc.add_heading(stripped, level=2)
            else:
                p = doc.add_paragraph(stripped)
                p.paragraph_format.space_after = Pt(6)

        safe_name = "".join(c for c in filename if c.isalnum() or c in " _-").strip()
        out_path = str(Path(self._export_path) / f"{safe_name}.docx")
        doc.save(out_path)
        logger.info("Exported DOCX: %s", out_path)
        return out_path


_service: "ExportService | None" = None


def get_export_service() -> ExportService:
    global _service
    if _service is None:
        _service = ExportService()
    return _service
